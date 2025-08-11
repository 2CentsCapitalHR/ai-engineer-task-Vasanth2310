# doc_parser.py
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def parse_docx(path):
    """
    Return list of paragraphs with indices and raw paragraph objects if needed.
    Output: [{"index": i, "text": text}, ...]
    """
    doc = Document(path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({"index": i, "text": text})
    return paragraphs

def annotate_docx(input_path, issues, output_path):
    """
    Create a copy of input_path with inline red comment-like annotations appended
    to the paragraph that was flagged.
    issues: list of dicts containing: paragraph_index, issue, severity, suggestion, citation, alt_clause (optional)
    The function appends an inline run (bold red) at the end of the flagged paragraph describing the issue.
    """
    in_doc = Document(input_path)
    out_doc = Document()

    # Build mapping from paragraph index to list of issues
    issues_by_idx = {}
    for it in issues:
        idx = it.get("paragraph_index")
        issues_by_idx.setdefault(idx, []).append(it)

    for idx, para in enumerate(in_doc.paragraphs):
        # copy original paragraph text preserving formatting at a simple level:
        p = out_doc.add_paragraph()
        # copy runs (best-effort): python-docx doesn't provide easy run copying; we'll set plain text
        p.add_run(para.text)

        # if there are issues for this paragraph, append inline comment-runs
        if idx in issues_by_idx:
            for issue in issues_by_idx[idx]:
                # create an annotation run
                ann = p.add_run(f"  ⚠️ ISSUE [{issue.get('severity','Medium')}]: {issue.get('issue')}")
                ann.bold = True
                ann.italic = False
                ann.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                # also append suggestion and citation on a new line within the paragraph (so it's easy to spot)
                sugg = p.add_run(f"\n    Suggestion: {issue.get('suggestion')}")
                sugg.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
                sugg.italic = True

                citation = issue.get('citation') or ""
                if citation:
                    cite_run = p.add_run(f"\n    Citation: {citation}")
                    cite_run.font.color.rgb = RGBColor(0x66, 0x00, 0x00)

                # optional alternative clause
                if issue.get('alt_clause'):
                    alt = p.add_run(f"\n    Alternative clause: {issue.get('alt_clause')}")
                    alt.italic = True
                    alt.font.color.rgb = RGBColor(0x33, 0x00, 0x00)

    out_doc.save(output_path)
